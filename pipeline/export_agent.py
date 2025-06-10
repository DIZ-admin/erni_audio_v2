# pipeline/export_agent.py

import datetime as _dt
import json
import logging
import time
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
from .base_agent import BaseAgent
from .validation_mixin import ValidationMixin


@dataclass
class ExportMetrics:
    """Метрики экспорта"""
    total_segments: int
    total_duration: float
    speakers_count: int
    total_words: int
    average_segment_duration: float
    export_formats: List[str]
    file_sizes: Dict[str, int]


class ExportAgent(BaseAgent, ValidationMixin):
    """
    Расширенный агент для экспорта в множественные форматы.

    Поддерживаемые форматы:
    - SRT: SubRip субтитры
    - JSON: Структурированные данные
    - ASS: Advanced SubStation Alpha
    - VTT: WebVTT субтитры
    - TTML: Timed Text Markup Language
    - TXT: Простой текст
    - CSV: Табличные данные
    - DOCX: Microsoft Word документ
    """

    SUPPORTED_FORMATS = ["srt", "json", "ass", "vtt", "ttml", "txt", "csv", "docx"]

    def __init__(self, format: str = "srt", create_all_formats: bool = False,
                 overwrite_existing: bool = False, add_timestamp: bool = False,
                 include_confidence: bool = False, speaker_colors: bool = True):
        """
        Args:
            format: Основной формат экспорта
            create_all_formats: Создать файлы во всех форматах
            overwrite_existing: Перезаписывать существующие файлы
            add_timestamp: Добавлять временную метку к именам файлов
            include_confidence: Включать информацию о уверенности
            speaker_colors: Использовать цвета для разных спикеров
        """
        # Инициализация базовых классов
        BaseAgent.__init__(self, name="ExportAgent")
        ValidationMixin.__init__(self)

        # Валидация формата через новый метод
        self.validate_export_format(format)

        self.format = format
        self.create_all_formats = create_all_formats
        self.overwrite_existing = overwrite_existing
        self.add_timestamp = add_timestamp
        self.include_confidence = include_confidence
        self.speaker_colors = speaker_colors

        # Цвета для спикеров (для форматов поддерживающих цвета)
        self.speaker_color_map = {
            "SPEAKER_00": "#FF6B6B",  # Красный
            "SPEAKER_01": "#4ECDC4",  # Бирюзовый
            "SPEAKER_02": "#45B7D1",  # Синий
            "SPEAKER_03": "#96CEB4",  # Зеленый
            "SPEAKER_04": "#FFEAA7",  # Желтый
            "SPEAKER_05": "#DDA0DD",  # Фиолетовый
            "SPEAKER_06": "#98D8C8",  # Мятный
            "SPEAKER_07": "#F7DC6F",  # Золотой
        }

        self.log_with_emoji("info", "📤", f"Инициализирован с форматом: {format}")
        if create_all_formats:
            self.log_with_emoji("info", "📋", f"Будут созданы все форматы: {', '.join(self.SUPPORTED_FORMATS)}")

    def validate_export_format(self, format: str) -> None:
        """
        Валидация формата экспорта.

        Args:
            format: Формат для валидации

        Raises:
            ValueError: Если формат не поддерживается
        """
        if not isinstance(format, str):
            raise ValueError(f"Формат должен быть строкой, получен {type(format)}")

        format = format.lower().strip()

        if not format:
            raise ValueError("Формат не может быть пустым")

        if format not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Неподдерживаемый формат: {format}. "
                f"Доступные: {', '.join(self.SUPPORTED_FORMATS)}"
            )

    def validate_export_segments(self, segments: List[Dict]) -> List[str]:
        """
        Валидация сегментов для экспорта.

        Args:
            segments: Список сегментов для валидации

        Returns:
            Список найденных проблем
        """
        issues = []

        if not isinstance(segments, list):
            issues.append("Сегменты должны быть списком")
            return issues

        if not segments:
            issues.append("Список сегментов пуст")
            return issues

        required_fields = ["start", "end", "speaker", "text"]

        for i, segment in enumerate(segments):
            if not isinstance(segment, dict):
                issues.append(f"Сегмент {i}: должен быть словарем")
                continue

            # Проверяем обязательные поля
            for field in required_fields:
                if field not in segment:
                    issues.append(f"Сегмент {i}: отсутствует поле '{field}'")

            # Проверяем типы данных
            if "start" in segment and not isinstance(segment["start"], (int, float)):
                issues.append(f"Сегмент {i}: 'start' должно быть числом")

            if "end" in segment and not isinstance(segment["end"], (int, float)):
                issues.append(f"Сегмент {i}: 'end' должно быть числом")

            if "speaker" in segment and not isinstance(segment["speaker"], str):
                issues.append(f"Сегмент {i}: 'speaker' должно быть строкой")

            if "text" in segment and not isinstance(segment["text"], str):
                issues.append(f"Сегмент {i}: 'text' должно быть строкой")

            # Проверяем временные метки
            if ("start" in segment and "end" in segment and
                isinstance(segment["start"], (int, float)) and
                isinstance(segment["end"], (int, float))):

                if segment["start"] < 0:
                    issues.append(f"Сегмент {i}: отрицательное время начала")

                if segment["end"] < 0:
                    issues.append(f"Сегмент {i}: отрицательное время окончания")

                if segment["start"] >= segment["end"]:
                    issues.append(f"Сегмент {i}: некорректные временные метки")

        return issues

    def validate_output_path(self, output_path: Path) -> None:
        """
        Валидация пути для экспорта.

        Args:
            output_path: Путь для валидации

        Raises:
            ValueError: Если путь невалиден
        """
        if not isinstance(output_path, Path):
            raise ValueError(f"output_path должен быть Path объектом, получен {type(output_path)}")

        # Проверяем, что родительская директория существует или может быть создана
        parent_dir = output_path.parent
        if not parent_dir.exists():
            try:
                parent_dir.mkdir(parents=True, exist_ok=True)
                self.log_with_emoji("info", "📁", f"Создана директория: {parent_dir}")
            except Exception as e:
                raise ValueError(f"Не удалось создать директорию {parent_dir}: {e}")

        # Проверяем права на запись
        if parent_dir.exists() and not parent_dir.is_dir():
            raise ValueError(f"Родительский путь не является директорией: {parent_dir}")

        # Проверяем, что файл не является директорией
        if output_path.exists() and output_path.is_dir():
            raise ValueError(f"Путь указывает на директорию, а не на файл: {output_path}")

    def get_speaker_color(self, speaker: str) -> str:
        """Возвращает цвет для спикера"""
        if not self.speaker_colors:
            return "#FFFFFF"  # Белый по умолчанию

        return self.speaker_color_map.get(speaker, "#CCCCCC")  # Серый для неизвестных

    def _ts_srt(self, sec: float) -> str:
        """Форматирование времени для SRT"""
        td = _dt.timedelta(seconds=sec)
        h, r = divmod(td.seconds, 3600)
        m, s = divmod(r, 60)
        ms = int(td.microseconds / 1000)
        return f"{h:02}:{m:02}:{s:02},{ms:03}"

    def _ts_ass(self, sec: float) -> str:
        """Форматирование времени для ASS"""
        td = _dt.timedelta(seconds=sec)
        h, r = divmod(td.seconds, 3600)
        m, s = divmod(r, 60)
        cs = int(td.microseconds / 10_000)  # centiseconds
        return f"{h:d}:{m:02}:{s:02}.{cs:02}"

    def _ts_vtt(self, sec: float) -> str:
        """Форматирование времени для VTT"""
        td = _dt.timedelta(seconds=sec)
        h, r = divmod(td.seconds, 3600)
        m, s = divmod(r, 60)
        ms = int(td.microseconds / 1000)
        return f"{h:02}:{m:02}:{s:02}.{ms:03}"

    def _ts_ttml(self, sec: float) -> str:
        """Форматирование времени для TTML"""
        return f"{sec:.3f}s"

    def write_srt(self, segs: List[Dict], outfile: Path):
        """Экспорт в SRT формат"""
        with open(outfile, "w", encoding="utf-8") as f:
            for idx, s in enumerate(segs, 1):
                f.write(f"{idx}\n{self._ts_srt(s['start'])} --> {self._ts_srt(s['end'])}\n")

                # Добавляем уверенность если включено
                confidence_info = ""
                if self.include_confidence and 'confidence' in s:
                    confidence_info = f" [{s['confidence']:.2f}]"

                f.write(f"{s['speaker']}: {s['text']}{confidence_info}\n\n")

    def write_json(self, segs: List[Dict], outfile: Path):
        """Экспорт в JSON формат"""
        # Добавляем метаданные
        export_data = {
            "metadata": {
                "export_timestamp": _dt.datetime.now().isoformat(),
                "total_segments": len(segs),
                "total_duration": max(s['end'] for s in segs) if segs else 0,
                "speakers": list(set(s['speaker'] for s in segs)),
                "include_confidence": self.include_confidence
            },
            "segments": segs
        }

        outfile.write_text(json.dumps(export_data, indent=2, ensure_ascii=False))

    def write_ass(self, segs: List[Dict], outfile: Path):
        """Экспорт в ASS формат с поддержкой цветов"""
        # Создаем стили для каждого спикера
        speakers = list(set(s['speaker'] for s in segs))
        styles = []

        for speaker in speakers:
            color = self.get_speaker_color(speaker)
            # Конвертируем hex в BGR формат для ASS
            color_bgr = f"&H00{color[5:7]}{color[3:5]}{color[1:3]}"

            styles.append(
                f"Style: {speaker},Arial,24,{color_bgr},&H000000FF,&H00000000,&H64000000,"
                f"-1,0,0,0,100,100,0,0,1,2,2,2,10,10,10,1"
            )

        header = (
            "[Script Info]\n"
            "Title: speech_pipeline export\n"
            "ScriptType: v4.00+\n"
            "Collisions: Normal\n"
            "PlayResX: 384\n"
            "PlayResY: 288\n"
            "[V4+ Styles]\n"
            "Format: Name,Fontname,Fontsize,PrimaryColour,SecondaryColour,OutlineColour,BackColour,"
            "Bold,Italic,Underline,StrikeOut,ScaleX,ScaleY,Spacing,Angle,BorderStyle,Outline,Shadow,"
            "Alignment,MarginL,MarginR,MarginV,Encoding\n"
        )

        # Добавляем стили
        header += "\n".join(styles) + "\n"

        header += (
            "[Events]\n"
            "Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text\n"
        )

        with open(outfile, "w", encoding="utf-8") as f:
            f.write(header)
            for s in segs:
                style = s['speaker'] if s['speaker'] in speakers else "Default"
                confidence_info = ""
                if self.include_confidence and 'confidence' in s:
                    confidence_info = f" [{s['confidence']:.2f}]"

                f.write(
                    f"Dialogue: 0,{self._ts_ass(s['start'])},{self._ts_ass(s['end'])},"
                    f"{style},{s['speaker']},0,0,0,,{s['text']}{confidence_info}\n"
                )

    def write_vtt(self, segs: List[Dict], outfile: Path):
        """Экспорт в WebVTT формат"""
        with open(outfile, "w", encoding="utf-8") as f:
            f.write("WEBVTT\n\n")

            for idx, s in enumerate(segs, 1):
                # WebVTT поддерживает цвета через CSS
                color = self.get_speaker_color(s['speaker']) if self.speaker_colors else None

                f.write(f"{idx}\n")
                f.write(f"{self._ts_vtt(s['start'])} --> {self._ts_vtt(s['end'])}\n")

                # Добавляем стиль если нужен цвет
                text = s['text']
                if color and self.speaker_colors:
                    text = f"<c.{s['speaker'].lower()}>{text}</c>"

                confidence_info = ""
                if self.include_confidence and 'confidence' in s:
                    confidence_info = f" [{s['confidence']:.2f}]"

                f.write(f"{s['speaker']}: {text}{confidence_info}\n\n")

    def write_ttml(self, segs: List[Dict], outfile: Path):
        """Экспорт в TTML формат"""
        # Создаем XML структуру
        root = ET.Element("tt")
        root.set("xmlns", "http://www.w3.org/ns/ttml")
        root.set("xmlns:tts", "http://www.w3.org/ns/ttml#styling")

        # Заголовок
        head = ET.SubElement(root, "head")
        styling = ET.SubElement(head, "styling")

        # Стили для спикеров
        speakers = list(set(s['speaker'] for s in segs))
        for speaker in speakers:
            style = ET.SubElement(styling, "style")
            style.set("xml:id", speaker.lower())
            if self.speaker_colors:
                style.set("tts:color", self.get_speaker_color(speaker))

        # Тело документа
        body = ET.SubElement(root, "body")
        div = ET.SubElement(body, "div")

        for s in segs:
            p = ET.SubElement(div, "p")
            p.set("begin", self._ts_ttml(s['start']))
            p.set("end", self._ts_ttml(s['end']))
            p.set("style", s['speaker'].lower())

            confidence_info = ""
            if self.include_confidence and 'confidence' in s:
                confidence_info = f" [{s['confidence']:.2f}]"

            p.text = f"{s['speaker']}: {s['text']}{confidence_info}"

        # Записываем XML
        tree = ET.ElementTree(root)
        tree.write(outfile, encoding="utf-8", xml_declaration=True)

    def write_txt(self, segs: List[Dict], outfile: Path):
        """Экспорт в простой текстовый формат"""
        with open(outfile, "w", encoding="utf-8") as f:
            # Заголовок
            f.write("ТРАНСКРИПЦИЯ АУДИО\n")
            f.write("=" * 50 + "\n")
            f.write(f"Дата экспорта: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Всего сегментов: {len(segs)}\n")
            if segs:
                f.write(f"Общая длительность: {max(s['end'] for s in segs):.2f} секунд\n")
            f.write("=" * 50 + "\n\n")

            current_speaker = None
            for s in segs:
                # Группируем по спикерам для лучшей читаемости
                if s['speaker'] != current_speaker:
                    if current_speaker is not None:
                        f.write("\n")
                    f.write(f"[{s['speaker']}]\n")
                    current_speaker = s['speaker']

                # Временные метки
                start_time = f"{int(s['start']//60):02d}:{int(s['start']%60):02d}"
                end_time = f"{int(s['end']//60):02d}:{int(s['end']%60):02d}"

                confidence_info = ""
                if self.include_confidence and 'confidence' in s:
                    confidence_info = f" (уверенность: {s['confidence']:.2f})"

                f.write(f"[{start_time}-{end_time}]{confidence_info} {s['text']}\n")

    def write_csv(self, segs: List[Dict], outfile: Path):
        """Экспорт в CSV формат"""
        import csv

        with open(outfile, "w", encoding="utf-8", newline='') as f:
            fieldnames = ['start_time', 'end_time', 'duration', 'speaker', 'text']
            if self.include_confidence:
                fieldnames.append('confidence')

            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()

            for s in segs:
                row = {
                    'start_time': f"{s['start']:.3f}",
                    'end_time': f"{s['end']:.3f}",
                    'duration': f"{s['end'] - s['start']:.3f}",
                    'speaker': s['speaker'],
                    'text': s['text']
                }

                if self.include_confidence and 'confidence' in s:
                    row['confidence'] = f"{s['confidence']:.3f}"

                writer.writerow(row)

    def write_docx(self, segs: List[Dict], outfile: Path):
        """Экспорт в DOCX формат"""
        try:
            from docx import Document
            from docx.shared import RGBColor
        except ImportError:
            self.log_with_emoji("error", "❌", "Для экспорта в DOCX требуется библиотека python-docx")
            self.log_with_emoji("info", "💡", "Установите: pip install python-docx")
            raise ImportError("python-docx не установлен")

        doc = Document()

        # Заголовок документа
        title = doc.add_heading('Транскрипция аудио', 0)

        # Метаданные
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Дата экспорта: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_para.add_run(f"Всего сегментов: {len(segs)}\n")
        if segs:
            meta_para.add_run(f"Общая длительность: {max(s['end'] for s in segs):.2f} секунд\n")

        doc.add_paragraph()  # Пустая строка

        # Группируем по спикерам
        current_speaker = None
        for s in segs:
            if s['speaker'] != current_speaker:
                # Новый спикер - добавляем заголовок
                speaker_heading = doc.add_heading(f"Спикер: {s['speaker']}", level=2)
                current_speaker = s['speaker']

            # Создаем параграф для сегмента
            para = doc.add_paragraph()

            # Временные метки
            start_time = f"{int(s['start']//60):02d}:{int(s['start']%60):02d}"
            end_time = f"{int(s['end']//60):02d}:{int(s['end']%60):02d}"

            time_run = para.add_run(f"[{start_time}-{end_time}] ")
            time_run.bold = True

            # Текст
            text_run = para.add_run(s['text'])

            # Цвет для спикера если включен
            if self.speaker_colors:
                color_hex = self.get_speaker_color(s['speaker'])
                rgb = tuple(int(color_hex[i:i+2], 16) for i in (1, 3, 5))
                text_run.font.color.rgb = RGBColor(*rgb)

            # Уверенность
            if self.include_confidence and 'confidence' in s:
                conf_run = para.add_run(f" (уверенность: {s['confidence']:.2f})")
                conf_run.italic = True

        doc.save(outfile)

    def _generate_unique_filename(self, output_path: Path, format: str) -> Path:
        """Генерирует уникальное имя файла"""
        corrected_path = self._ensure_correct_extension(output_path, format)

        if self.overwrite_existing:
            return corrected_path

        if not corrected_path.exists():
            return corrected_path

        base_path = corrected_path.with_suffix('')
        extension = corrected_path.suffix

        if self.add_timestamp:
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            unique_path = Path(f"{base_path}_{timestamp}{extension}")

            counter = 1
            while unique_path.exists():
                unique_path = Path(f"{base_path}_{timestamp}_{counter:03d}{extension}")
                counter += 1
        else:
            counter = 1
            unique_path = Path(f"{base_path}_{counter:03d}{extension}")
            while unique_path.exists():
                counter += 1
                unique_path = Path(f"{base_path}_{counter:03d}{extension}")

        self.logger.info(f"📝 Файл {corrected_path} уже существует, создаю уникальное имя: {unique_path}")
        return unique_path



    def _export_single_format(self, merged: List[Dict], output_path: Path, format: str):
        """Экспортирует данные в один конкретный формат"""
        export_methods = {
            "srt": self.write_srt,
            "json": self.write_json,
            "ass": self.write_ass,
            "vtt": self.write_vtt,
            "ttml": self.write_ttml,
            "txt": self.write_txt,
            "csv": self.write_csv,
            "docx": self.write_docx
        }

        if format not in export_methods:
            raise ValueError(f"Неподдерживаемый формат: {format}")

        try:
            export_methods[format](merged, output_path)
            self.log_with_emoji("info", "✅", f"Экспорт в {format.upper()} завершен: {output_path}")
        except Exception as e:
            self.log_with_emoji("error", "❌", f"Ошибка экспорта в {format.upper()}: {e}")
            raise

    def calculate_export_metrics(self, merged: List[Dict], created_files: List[Path]) -> ExportMetrics:
        """Вычисляет метрики экспорта"""
        if not merged:
            return ExportMetrics(
                total_segments=0, total_duration=0.0, speakers_count=0,
                total_words=0, average_segment_duration=0.0,
                export_formats=[], file_sizes={}
            )

        # Базовые метрики
        total_segments = len(merged)
        total_duration = max(s['end'] for s in merged) if merged else 0
        speakers = set(s['speaker'] for s in merged)
        speakers_count = len(speakers)

        # Подсчет слов
        total_words = sum(len(s['text'].split()) for s in merged)

        # Средняя длительность сегмента
        segment_durations = [s['end'] - s['start'] for s in merged]
        average_segment_duration = sum(segment_durations) / len(segment_durations)

        # Размеры файлов
        file_sizes = {}
        export_formats = []

        for file_path in created_files:
            if file_path.exists():
                format_name = file_path.suffix[1:].upper()  # Убираем точку
                file_sizes[format_name] = file_path.stat().st_size
                export_formats.append(format_name)

        return ExportMetrics(
            total_segments=total_segments,
            total_duration=total_duration,
            speakers_count=speakers_count,
            total_words=total_words,
            average_segment_duration=average_segment_duration,
            export_formats=export_formats,
            file_sizes=file_sizes
        )

    def run(self, merged: List[Dict], output_path: Path) -> List[Path]:
        """
        Основной метод экспорта с расширенной функциональностью

        Args:
            merged: Объединенные сегменты для экспорта
            output_path: Путь для сохранения файлов

        Returns:
            Список созданных файлов
        """
        self.start_operation("экспорт")

        try:
            # Валидация входных данных
            segment_issues = self.validate_export_segments(merged)
            if segment_issues:
                self.log_with_emoji("warning", "⚠️", f"Проблемы в сегментах: {len(segment_issues)}")
                for issue in segment_issues[:3]:  # Показываем первые 3
                    self.log_with_emoji("warning", "   ", issue)

                if not merged:  # Если сегменты пусты, завершаем
                    self.end_operation("экспорт", success=True)
                    return []

            # Валидация пути вывода
            self.validate_output_path(output_path)

            self.log_with_emoji("info", "📤", f"Начинаю экспорт {len(merged)} сегментов...")

            created_files = []

            if self.create_all_formats:
                # Создаем файлы во всех форматах
                base_path = output_path.with_suffix('')

                for fmt in self.SUPPORTED_FORMATS:
                    try:
                        fmt_path = self._generate_unique_filename(base_path, fmt)
                        self._export_single_format(merged, fmt_path, fmt)
                        created_files.append(fmt_path)
                    except Exception as e:
                        self.log_with_emoji("error", "❌", f"Ошибка экспорта в {fmt.upper()}: {e}")
                        # Продолжаем с другими форматами

                self.log_with_emoji("info", "✅", f"Создано файлов: {len(created_files)} из {len(self.SUPPORTED_FORMATS)} форматов")
            else:
                # Создаем файл только в указанном формате
                unique_path = self._generate_unique_filename(output_path, self.format)
                self._export_single_format(merged, unique_path, self.format)
                created_files.append(unique_path)

                self.log_with_emoji("info", "✅", f"Создан файл в формате {self.format.upper()}: {unique_path}")

            # Вычисляем и логируем метрики
            if created_files:
                metrics = self.calculate_export_metrics(merged, created_files)
                self.log_with_emoji("info", "📊", "Метрики экспорта:")
                self.log_with_emoji("info", "   ", f"Сегментов: {metrics.total_segments}")
                self.log_with_emoji("info", "   ", f"Спикеров: {metrics.speakers_count}")
                self.log_with_emoji("info", "   ", f"Слов: {metrics.total_words}")
                self.log_with_emoji("info", "   ", f"Длительность: {metrics.total_duration:.2f}с")

                # Размеры файлов
                for fmt, size in metrics.file_sizes.items():
                    size_kb = size / 1024
                    self.log_with_emoji("info", "   ", f"{fmt}: {size_kb:.1f} KB")

            self.end_operation("экспорт", success=True)
            return created_files

        except Exception as e:
            self.end_operation("экспорт", success=False)
            self.handle_error(e, "экспорт", reraise=True)

    def _ts_srt(self, sec: float) -> str:
        td = _dt.timedelta(seconds=sec)
        h, r = divmod(td.seconds, 3600)
        m, s = divmod(r, 60)
        ms = int(td.microseconds / 1000)
        return f"{h:02}:{m:02}:{s:02},{ms:03}"

    def _ts_ass(self, sec: float) -> str:
        td = _dt.timedelta(seconds=sec)
        h, r = divmod(td.seconds, 3600)
        m, s = divmod(r, 60)
        cs = int(td.microseconds / 10_000)  # centiseconds
        return f"{h:d}:{m:02}:{s:02}.{cs:02}"

    def _generate_unique_filename(self, output_path: Path, format: str) -> Path:
        """
        Генерирует уникальное имя файла, избегая перезаписи существующих файлов.
        Добавляет временную метку или числовой суффикс при необходимости.
        """
        corrected_path = self._ensure_correct_extension(output_path, format)

        # Если перезапись разрешена, возвращаем исходный путь
        if self.overwrite_existing:
            return corrected_path

        # Если файл не существует, возвращаем исходный путь
        if not corrected_path.exists():
            return corrected_path

        # Генерируем уникальное имя
        base_path = corrected_path.with_suffix('')
        extension = corrected_path.suffix

        if self.add_timestamp:
            # Добавляем временную метку
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            unique_path = Path(f"{base_path}_{timestamp}{extension}")

            # Если файл с временной меткой тоже существует, добавляем счетчик
            counter = 1
            while unique_path.exists():
                unique_path = Path(f"{base_path}_{timestamp}_{counter:03d}{extension}")
                counter += 1
        else:
            # Добавляем числовой суффикс
            counter = 1
            unique_path = Path(f"{base_path}_{counter:03d}{extension}")
            while unique_path.exists():
                counter += 1
                unique_path = Path(f"{base_path}_{counter:03d}{extension}")

        self.log_with_emoji("info", "📝", f"Файл {corrected_path} уже существует, создаю уникальное имя: {unique_path}")
        return unique_path

    def _ensure_correct_extension(self, output_path: Path, format: str) -> Path:
        """
        Обеспечивает правильное расширение файла для указанного формата.
        Если файл уже имеет правильное расширение, возвращает его как есть.
        Если нет - добавляет правильное расширение.
        """
        format_extensions = {
            "srt": ".srt",
            "json": ".json",
            "ass": ".ass",
            "vtt": ".vtt",
            "ttml": ".ttml",
            "txt": ".txt",
            "csv": ".csv",
            "docx": ".docx"
        }

        if format not in format_extensions:
            raise ValueError(f"Неподдерживаемый формат: {format}")

        expected_ext = format_extensions[format]

        # Если файл уже имеет правильное расширение
        if output_path.suffix.lower() == expected_ext:
            return output_path

        # Если файл имеет другое расширение или не имеет расширения
        if output_path.suffix:
            # Заменяем существующее расширение
            return output_path.with_suffix(expected_ext)
        else:
            # Добавляем расширение к файлу без расширения
            return Path(str(output_path) + expected_ext)

    def write_json(self, segs: List[Dict], outfile: Path):
        outfile.write_text(json.dumps(segs, indent=2, ensure_ascii=False))

    def write_srt(self, segs: List[Dict], outfile: Path):
        with open(outfile, "w", encoding="utf-8") as f:
            for idx, s in enumerate(segs, 1):
                f.write(f"{idx}\n{self._ts_srt(s['start'])} --> {self._ts_srt(s['end'])}\n")
                f.write(f"{s['speaker']}: {s['text']}\n\n")

    def write_ass(self, segs: List[Dict], outfile: Path):
        header = (
            "[Script Info]\n"
            "Title: speech_pipeline export\n"
            "ScriptType: v4.00+\n"
            "Collisions: Normal\n"
            "PlayResX: 384\n"
            "PlayResY: 288\n"
            "[V4+ Styles]\n"
            "Format: Name,Fontname,Fontsize,PrimaryColour,SecondaryColour,OutlineColour,BackColour,"
            "Bold,Italic,Underline,StrikeOut,ScaleX,ScaleY,Spacing,Angle,BorderStyle,Outline,Shadow,"
            "Alignment,MarginL,MarginR,MarginV,Encoding\n"
            "Style: Default,Arial,24,&H00FFFFFF,&H000000FF,&H00000000,&H64000000,"
            "-1,0,0,0,100,100,0,0,1,2,2,2,10,10,10,1\n"
            "[Events]\n"
            "Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text\n"
        )
        with open(outfile, "w", encoding="utf-8") as f:
            f.write(header)
            for s in segs:
                f.write(
                    f"Dialogue: 0,{self._ts_ass(s['start'])},{self._ts_ass(s['end'])},"
                    f"Default,{s['speaker']},0,0,0,,{s['text']}\n"
                )

    # Дублированный метод run удален - используется основной метод выше
